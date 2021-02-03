from docx import Document
from docx.shared import Inches
import os

input_path = r'C:\Users\adam.smith04\Documents\PROJECTS\mu_code\wordautomation\resources\test.docx'
output_path = r'C:\Users\adam.smith04\Documents\PROJECTS\mu_code\wordautomation\resources\Output.docx'
doc = Document(input_path)



table = doc.tables[0]


cell1_header = table.cell(0,0)
cell2_header = table.cell(1,0)
cell3_header = table.cell(2,0)
cell4_header = table.cell(3,0)
cell5_header = table.cell(4,0)
cell6_header = table.cell(5,0)
cell7_header = table.cell(6,0)
cell8_header = table.cell(7,0)


cell1_value = table.cell(0,1)
cell2_value = table.cell(1,1)
cell3_value = table.cell(2,1)
cell4_value = table.cell(3,1)
cell5_value = table.cell(4,1)
cell6_value = table.cell(5,1)
cell7_value = table.cell(6,1)
cell8_value = table.cell(7,1)



print('inputing data...')
cell1_value.add_paragraph(text="Dr. Young")
cell2_value.add_paragraph(text="HCP Peeps")
cell3_value.add_paragraph(text="555-555-5555")
cell4_value.add_paragraph(text="Bronte Woodlan")
cell5_value.add_paragraph(text="04-02-1994")
cell6_value.add_paragraph(text="42069")
cell7_value.add_paragraph(text="Adam Smith")
cell8_value.add_paragraph(text="02/03/2021")

print('saving new file...')
doc.save(output_path)
print('opening file...')
os.startfile(output_path)